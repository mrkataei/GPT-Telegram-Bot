from telegram import *
import docx
import PyPDF2
from pydub import AudioSegment
from sqlalchemy.orm import Session
from time import sleep
from concurrent.futures import ThreadPoolExecutor

from crud.user import user as crudUser
from crud.openai import API as crudAPI
from schema.user import UserCreate
from models.subscribe import Subscribe
from models.openai import StatusEnum, OpenAI as modelAI

bot = TeleBot(BOT_TOKEN)


def is_commands(message) -> bool:
    commands = ['/contact', '/help', '/charge', trans("C_OCR"), trans("C_contact"),
                trans("C_Help"), trans("C_buy_token"), trans(
                    "C_backhome"), trans("C_invite_code"),
                '/start', trans("C_billings"), trans(
                    "C_profile_status"), trans("C_Transcript"),
                trans("C_translation"), trans("C_change_language"),
                trans("C_summarize"), trans(
                    "C_exchange"), trans("C_youtube_summarize"),
                trans("C_generate_picture")]

    if message.text not in commands or message.content_type != 'text':
        return False
    else:
        return True


def delete_extra_file(path: str) -> bool:
    if os.path.exists(path):
        os.remove(path)
        return True
    else:
        return False


def try_again_exception(error_name: str):
    def decorator(fun):
        def wrapper(message, *args, **kwg):
            try:
                fun(message, *args, **kwg)
            except Exception as error:
                logger.error(
                    f"{datetime.now()}: {error_name} for user : {message.chat.id}. error is :\n{error}")
                try:
                    bot.send_message(message.chat.id, trans("M_try_again"),
                                     reply_markup=start_keyboard())
                except Exception as sending:
                    logger.info(
                        f"{datetime.now()}: {error_name} for user : {message.chat.id}. blocked {sending}")

        return wrapper
    return decorator


def generate_invite_code() -> Tuple[str, str]:
    code = ''.join(random.choices('0123456789ABCDEF', k=8))
    link = f"https://t.me/GPT_fa_bot?start={code}"
    return code, link


def success_invite(username_invited: str, first_name_invited: str, chat_id: str):
    try:
        bot.send_message(chat_id=chat_id,
                         text=trans("M_invited").format(username_invited=username_invited,
                                                        first_name_invited=first_name_invited))
    except Exception as e:
        logger.error(
            f"{datetime.now()}: success invite for user : {chat_id}. error is :\n{e}")


def ocr(file_path: str, chat_id: int) -> bool:
    try:
        with open(file_path, 'rb') as pdf_file:
            # Create a PDF reader object
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            # Create a new Word document
            word_document = docx.Document()

            # Iterate through each page in the PDF
            for page in pdf_reader.pages:
                # Extract the text from the page
                text = page.extract_text()
                # Add a new paragraph to the Word document with the extracted text
                word_document.add_paragraph(text)

            output_file = f'{file_path}_output.docx'
            # Save the Word document to a file
            word_document.save(output_file)

        bot.send_document(chat_id, document=open(
            output_file, 'rb'), reply_markup=START_KEYBOARD)

        os.remove(file_path)
        os.remove(output_file)

        return True
    except Exception as e:
        logger.error(
            f"{datetime.now()}: ocr for user : {chat_id}. error is :\n{e}")
        return False


def convert_ogg_to_mp3(path_ogg: str) -> str:
    try:
        ogg_file = AudioSegment.from_file(path_ogg, format="ogg")
        # Convert to MP3 format
        delete_extra_file(path=path_ogg)
        output = f"{path_ogg}.mp3"
        ogg_file.export(output, format="mp3")
        return output
    except Exception as e:
        logger.error(f"{datetime.now()}: convert_ogg_to_mp3 . error is :\n{e}")
        return None


def read_docx(db: Session, chat_id: int, file_path: str, lang: str = 'english') -> str:
    try:
        doc = docx.Document(file_path)
        text_output = ""
        for paragraph in doc.paragraphs:
            text_output += "\n" + paragraph.text
        text_output = f"translate this to {lang}\n\n" + text_output
        if len(str(text_output).split()) > MESSAGE_SIZE_LIMIT:
            bot.send_message(chat_id, trans("M_out_of_size_message"),
                             reply_markup=START_KEYBOARD)
            os.remove(file_path)
            return False
        try:
            completion = openai.ChatCompletion.create(
                model=MODEL, messages=[{"role": "user", "content": text_output}])

        except openai.OpenAIError as e:
            
            if "Rate limit" in str(e):
                logger.warning(
                    f"{datetime.now()}: translate docx - read docx {chat_id}. error is :\n{e}")
                change_api_status(db=db, api=openai.api_key, status=StatusEnum.PENDING)
            
            elif "billing" in str(e) or "exceeded" in str(e):
                logger.info(
                        f"{datetime.now()}: chat handler single message {message.chat.id}. error is :Billing hard limit has been reached")
                change_api_status(db=db, api=openai.api_key, status=StatusEnum.INACTIVE) 
                
            os.remove(file_path)
            return False

        result = completion.choices[0].message.content
        result_output = docx.Document()
        result_output.add_paragraph(result)
        output_file = f'{file_path}_output.docx'

        # Save the Word document to a file
        result_output.save(output_file)
        bot.send_document(chat_id, document=open(
            output_file, 'rb'), reply_markup=START_KEYBOARD)

        os.remove(file_path)
        os.remove(output_file)
        return True
    except Exception as e:
        logger.error(
            f"{datetime.now()}: read docx for user : {chat_id}. error is :\n{e}")
        os.remove(file_path)
        return False


def is_valid_id(chat_id: int) -> bool:
    return True if chat_id > 0 else False


def start_bot(chat_id: str, db: Session, username: str = "", first_name: str = "", is_start: bool = True, referral: str = ""):
    code, link = generate_invite_code()
    try:
        user_obj_in = UserCreate(chat_id=chat_id, username=username, first_name=first_name,
                                 referral_code=code, invite_link=link)
        crudUser.create(db=db, obj_in=user_obj_in)

        if is_start:
            user_inv = crudUser.get_by_referral(db=db, referral_code=referral)
            if user_inv:
                crudUser.charge_token(db=db, chat_id=user_inv.chat_id,
                                      value=BONUS)
                success_invite(username_invited=username, first_name_invited=first_name,
                               chat_id=user_inv.chat_id)

    except Exception as e:
        try:
            bot.send_message(chat_id, trans("M_try_again"),
                             reply_markup=START_KEYBOARD)
        except Exception as t:
            logger.warning(
                f"{datetime.now()}: sending message to user {chat_id} warning is {t}")
        logger.error(
            f"{datetime.now()}: chack subscribe for user {chat_id}. error is :\n{e}")


def check_subscribe(chat_id: str, channels: List[Subscribe]) -> bool:
    keyboard = join_keyboard(channels=channels)
    result = True
    for channel in channels:
        try:
            chat_member = bot.get_chat_member(chat_id=channel.channel_id,
                                              user_id=chat_id)
            if chat_member.status not in ['member', 'creator', 'administrator']:
                result = False
                bot.send_message(chat_id=chat_id, text=trans("M_join_first"),
                                 reply_markup=keyboard)
        except apihelper.ApiException as e:
            logger.error(
                f"{datetime.now()}: chack subscribe for user {chat_id}. error is :\n{e}")
            return False
    return result
        


def command_validate(message, db: Session, channels: List[Subscribe], token: int = 0) -> bool:
    chat_id = str(message.chat.id)
    if not is_valid_id(int(chat_id)):
        return False

    user = crudUser.get_by_chat_id(db=db, chat_id=chat_id)
    if not user:
        username = message.chat.username if message.chat.username else ""
        first_name = str(message.chat.first_name).split(" ")[
            0] if message.chat.first_name else ""
        refferal = ""
        is_start = False
        if message.text.startswith('/start '):
            refferal = message.text.split()[1]
            is_start = True
        start_bot(chat_id=chat_id, db=db, username=username,
                  first_name=first_name, is_start=is_start, referral=refferal)

    elif not crudUser.has_token(db_oj=user, token=token):
        keyboard = InlineKeyboardMarkup(row_width=1)
        keyboard.add(InlineKeyboardButton(text=trans("C_buy_token"),
                                          callback_data="buy_token"))
        bot.send_message(chat_id=chat_id,
                         text=trans("M_charge_your_account"),
                         reply_markup=keyboard)
        return False

    return check_subscribe(chat_id=chat_id, channels=channels)


def back_home(message: Dict, func) -> bool:
    if message.text == trans("C_backhome"):
        func(message=message)
        return True
    return False


def send_waiting_action(chat_id: str, token_usage: int, score: int,
                        need_action: bool = True, action: str = "typing") -> None:
    bot.send_message(chat_id, trans("M_waiting").format(token=token_usage,
                                                        score=score),
                     parse_mode="MarkdownV2")
    if need_action:
        bot.send_chat_action(chat_id=chat_id, action=action)


def change_api_status(db: Session, api: str, status: StatusEnum):
    obj = crudAPI.get(db=db, api=api)
    if obj:
        crudAPI.update_status(db=db, db_obj=obj, status=status)
        new_api = crudAPI.get_first_active(db=db)
        openai.api_key = new_api.api
        logger.info(
            f"{datetime.now()}: API successfully changed to {new_api.api} for email {new_api.email} reason: {status.name}")
        if status == StatusEnum.PENDING:
            pool = ThreadPoolExecutor(4)
            pool.submit(waiting_for_limited_apis, db=db, obj=obj)


def waiting_for_limited_apis(db: Session, obj: modelAI):
    logger.info(
        f"{datetime.now()}: API {obj.api} for email {obj.email} waiting for activating")
    sleep(25)
    logger.info(
        f"{datetime.now()}: API {obj.api} for email {obj.email} is now active")
    crudAPI.update_status(db=db, db_obj=obj, status=StatusEnum.ACTIVE)
